"""report_writer.py - JSON/TXT/HTML/DOCX report generation."""
import json
import html as html_lib
from collections import Counter
from datetime import datetime, timezone

SEVERITY_ORDER = ["Critical", "High", "Medium", "Low", "Info"]
SEV_WEIGHT = {s: i for i, s in enumerate(SEVERITY_ORDER)}


def sort_findings(findings: list) -> list:
    return sorted(findings, key=lambda f: SEV_WEIGHT.get(f.get("severity", "Info"), 99))


def severity_summary(findings: list) -> dict:
    counts = Counter(f.get("severity", "Info") for f in findings)
    return {s: counts.get(s, 0) for s in SEVERITY_ORDER}


def _now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


# ── JSON ───────────────────────────────────────────────────────────────────────

def write_json(path: str, target: str, findings: list, stats: dict, recon: dict = None):
    data = {
        "generated": _now(),
        "target": target,
        "stats": stats,
        "severity_summary": severity_summary(findings),
        "findings": findings,
        "recon": recon or {},
    }
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, default=str)


# ── TXT ────────────────────────────────────────────────────────────────────────

def write_txt(path: str, target: str, findings: list, stats: dict):
    lines = [
        "=" * 70,
        "  WebVulnScanner — Vulnerability Report",
        f"  Generated : {_now()}",
        f"  Target    : {target}",
        f"  Pages     : {stats.get('pages_crawled', '?')}",
        f"  Findings  : {stats.get('total_findings', '?')}",
        "=" * 70, "",
    ]
    sev = severity_summary(findings)
    lines.append("SEVERITY BREAKDOWN")
    for s in SEVERITY_ORDER:
        if sev[s]:
            lines.append(f"  {s:10}: {sev[s]}")
    lines += ["", "-" * 70, "DETAILED FINDINGS", ""]

    for i, f in enumerate(findings, 1):
        lines += [
            f"[{i}] {f.get('type', 'Unknown')}",
            f"    Severity    : {f.get('severity', '?')}",
            f"    URL         : {f.get('url', '?')}",
            f"    Location    : {f.get('location', '?')}",
            f"    Description : {f.get('description', '')}",
            f"    Evidence    : {f.get('evidence', '')}",
            f"    Remediation : {f.get('recommendation', '')}",
            "",
        ]
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))


# ── HTML ───────────────────────────────────────────────────────────────────────

SEV_COLORS = {
    "Critical": "#ff4444", "High": "#ff8c00",
    "Medium":   "#ffc107", "Low": "#17c964", "Info": "#6b7280",
}


def write_html(path: str, target: str, findings: list, stats: dict, recon: dict = None):
    e = html_lib.escape
    sev = severity_summary(findings)

    rows = ""
    for i, f in enumerate(findings, 1):
        color = SEV_COLORS.get(f.get("severity", "Info"), "#888")
        rows += f"""
        <tr>
          <td>{i}</td>
          <td>{e(f.get('type',''))}</td>
          <td style="color:{color};font-weight:bold">{e(f.get('severity',''))}</td>
          <td>{e(f.get('url',''))}</td>
          <td>{e(f.get('location',''))}</td>
          <td>{e(f.get('description',''))}</td>
          <td><code>{e(f.get('evidence',''))}</code></td>
          <td>{e(f.get('recommendation',''))}</td>
        </tr>"""

    sev_badges = " ".join(
        f'<span style="background:{SEV_COLORS.get(s,"#888")};color:#000;'
        f'padding:3px 8px;border-radius:4px;margin-right:6px">'
        f'{s}: {sev[s]}</span>'
        for s in SEVERITY_ORDER if sev[s]
    )

    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>WebVulnScanner Report — {e(target)}</title>
  <style>
    body{{font-family:monospace;background:#0a0f0a;color:#00ff41;margin:0;padding:2rem}}
    h1{{color:#00ff41;border-bottom:1px solid #00ff4140;padding-bottom:.5rem}}
    table{{width:100%;border-collapse:collapse;margin-top:1rem;font-size:13px}}
    th{{background:#112211;color:#00ff41;padding:8px;text-align:left;border:1px solid #1a2a1a}}
    td{{padding:7px 8px;border:1px solid #1a2a1a;vertical-align:top}}
    tr:nth-child(even){{background:#0d160d}}
    code{{color:#7dff7d;word-break:break-all}}
    .meta{{color:#5a8a5a;font-size:13px;margin-bottom:1.5rem}}
  </style>
</head>
<body>
  <h1>&#9632; WebVulnScanner Report</h1>
  <div class="meta">
    Target: <strong>{e(target)}</strong> &nbsp;|&nbsp;
    Generated: {_now()} &nbsp;|&nbsp;
    Pages: {stats.get('pages_crawled','?')} &nbsp;|&nbsp;
    Findings: {stats.get('total_findings','?')}
  </div>
  <div style="margin-bottom:1.5rem">{sev_badges}</div>
  <table>
    <thead><tr>
      <th>#</th><th>Type</th><th>Severity</th><th>URL</th>
      <th>Location</th><th>Description</th><th>Evidence</th><th>Remediation</th>
    </tr></thead>
    <tbody>{rows}</tbody>
  </table>
</body>
</html>"""
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(html_content)


# ── DOCX ───────────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    return "".join(c for c in str(text) if ord(c) >= 32 or c in "\n\r\t")


def write_docx(path: str, target: str, findings: list, stats: dict):
    from docx import Document
    from docx.shared import RGBColor

    doc = Document()
    doc.add_heading("WebVulnScanner Vulnerability Report", 0)

    p = doc.add_paragraph()
    p.add_run("Target: ").bold = True;   p.add_run(_clean(target))
    p = doc.add_paragraph()
    p.add_run("Generated: ").bold = True; p.add_run(_now())
    p = doc.add_paragraph()
    p.add_run("Total Findings: ").bold = True
    p.add_run(str(stats.get("total_findings", "?")))

    doc.add_heading("Severity Breakdown", level=2)
    sev = severity_summary(findings)
    tbl = doc.add_table(rows=1, cols=len([s for s in SEVERITY_ORDER if sev[s]]))
    tbl.style = "Table Grid"
    for i, s in enumerate(s for s in SEVERITY_ORDER if sev[s]):
        tbl.cell(0, i).text = _clean(f"{s}: {sev[s]}")

    doc.add_heading("Detailed Findings", level=1)
    if not findings:
        doc.add_paragraph("No vulnerabilities identified.")
    else:
        for i, f in enumerate(findings, 1):
            doc.add_heading(_clean(f"{i}. {f.get('type','')}"), level=2)
            p = doc.add_paragraph(); p.add_run("Severity: ").bold = True
            sr = p.add_run(_clean(f.get("severity", "")))
            if f.get("severity") == "Critical":
                sr.font.color.rgb = RGBColor(255, 0, 0)
            elif f.get("severity") == "High":
                sr.font.color.rgb = RGBColor(255, 140, 0)
            for label, key in [("URL", "url"), ("Location", "location")]:
                p = doc.add_paragraph(); p.add_run(f"{label}: ").bold = True
                p.add_run(_clean(f.get(key, "")))
            doc.add_heading("Description", level=3)
            doc.add_paragraph(_clean(f.get("description", "")))
            doc.add_heading("Evidence", level=3)
            doc.add_paragraph(_clean(f.get("evidence", "n/a")), style="No Spacing")
            doc.add_heading("Remediation", level=3)
            doc.add_paragraph(_clean(f.get("recommendation", "n/a")))
            doc.add_page_break()

    doc.save(path)
